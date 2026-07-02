from __future__ import annotations

from pathlib import Path
import textwrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


OUTPUT_DIR = Path(__file__).resolve().parent
BASENAME = "sample_procurement_contract"

TITLE = "采购服务合同测试样本"
SUBTITLE = "用于 Apache Tika 文件解析功能测试"

CONTRACT_LINES = [
    "合同编号：CTR-PARSE-2026-001",
    "合同类型：采购合同",
    "甲方：上海星河智能设备有限公司",
    "乙方：杭州云启技术服务有限公司",
    "合同金额：人民币565,000.00元",
    "不含税金额：人民币500,000.00元",
    "税率：13%",
    "签订日期：2026年04月16日",
    "生效日期：2026年04月16日",
    "有效期至：2027年04月15日",
    "履约地点：上海市浦东新区张江园区",
    "业务主办部门：采购部",
]

CLAUSES = [
    (
        "第一条 合同标的",
        "甲方向乙方采购合同智能审查平台部署及年度运维服务。乙方应向甲方提供系统初始化、权限配置、数据导入、用户培训和一年期运维支持。乙方交付内容包括部署包、管理员手册、验收报告和不少于两次现场培训。",
    ),
    (
        "第二条 交付与验收",
        "乙方应于合同生效后30个自然日内完成系统部署，并在甲方指定环境完成联调。甲方应在收到乙方交付物后10个工作日内组织验收。若验收发现重大缺陷，乙方应在5个工作日内完成整改并重新提交验收。",
    ),
    (
        "第三条 付款条件",
        "甲方在收到乙方开具的合法有效增值税专用发票、验收证明和付款申请后30个自然日内支付合同含税金额的70%；系统稳定运行满90日且无重大缺陷后，甲方支付剩余30%。付款方式为银行转账。",
    ),
    (
        "第四条 发票与税务",
        "乙方应按照13%的增值税税率向甲方开具增值税专用发票。若因乙方发票信息错误、发票逾期或不符合税务规定导致甲方无法抵扣进项税，乙方应负责重新开票并承担因此产生的损失。",
    ),
    (
        "第五条 违约责任",
        "乙方逾期交付的，每逾期一日应按合同含税总金额的万分之五向甲方支付违约金；逾期超过15日的，甲方有权解除合同并要求乙方退还已收款项。甲方逾期付款的，应按逾期未付款金额的同期一年期贷款市场报价利率承担资金占用损失。",
    ),
    (
        "第六条 保密与数据安全",
        "双方应对在履约过程中获得的合同、价格、业务数据、系统账号和技术资料承担保密义务。未经对方书面同意，任何一方不得向第三方披露。乙方接触甲方数据时，应采取访问控制、日志留存和数据脱敏措施。",
    ),
    (
        "第七条 关联交易与廉洁合规",
        "乙方承诺其与甲方采购经办人员不存在未披露的关联关系或利益冲突。乙方不得向甲方人员提供回扣、礼品、旅游或其他不正当利益。若发现关联交易或廉洁风险，甲方有权暂停付款并提交合规负责人复核。",
    ),
    (
        "第八条 争议解决",
        "因本合同产生的争议，双方应先友好协商解决；协商不成的，任何一方均可向甲方所在地有管辖权的人民法院提起诉讼。争议期间，除争议事项外，双方应继续履行未受影响的合同义务。",
    ),
    (
        "第九条 合同生效与期限",
        "本合同自双方法定代表人或授权代表签字并加盖公章之日起生效，有效期至2027年04月15日。合同期满后，如双方继续合作，应另行签署书面协议。",
    ),
    (
        "第十条 附件与补充材料",
        "本合同附件包括报价单、项目实施计划、验收标准、付款审批单模板和售后服务承诺书。附件与本合同具有同等法律效力；附件与正文不一致的，以正文约定为准。",
    ),
]

SIGNATURE_LINES = [
    "甲方（盖章）：上海星河智能设备有限公司",
    "授权代表：______________",
    "乙方（盖章）：杭州云启技术服务有限公司",
    "授权代表：______________",
]


def contract_text() -> str:
    parts = [TITLE, SUBTITLE, "", *CONTRACT_LINES, ""]
    for title, body in CLAUSES:
        parts.extend([title, body, ""])
    parts.extend(SIGNATURE_LINES)
    return "\n".join(parts).strip() + "\n"


def set_run_font(run, font_name: str, size: int | None = None, color: str | None = None, bold: bool = False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def set_style_font(style, font_name: str, size: int, color: str = "000000", bold: bool = False):
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def write_txt(path: Path):
    path.write_text(contract_text(), encoding="utf-8")


def write_docx(path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    set_style_font(doc.styles["Normal"], "Microsoft YaHei", 11)
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.10
    set_style_font(doc.styles["Heading 1"], "Microsoft YaHei", 16, "2E74B5", True)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(16)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(8)
    set_style_font(doc.styles["Heading 2"], "Microsoft YaHei", 13, "2E74B5", True)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(12)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(TITLE)
    set_run_font(run, "Microsoft YaHei", 18, "0B2545", True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run(SUBTITLE)
    set_run_font(run, "Microsoft YaHei", 10, "555555")

    doc.add_heading("合同主数据", level=1)
    for line in CONTRACT_LINES:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(3)

    doc.add_heading("合同条款", level=1)
    for title_text, body in CLAUSES:
        doc.add_heading(title_text, level=2)
        for paragraph in textwrap.wrap(body, width=88):
            doc.add_paragraph(paragraph)

    doc.add_heading("签署页", level=1)
    for line in SIGNATURE_LINES:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(8)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("解析测试样本 - 非真实合同")
    set_run_font(run, "Microsoft YaHei", 9, "777777")

    doc.save(path)


def register_pdf_font() -> str:
    font_candidates = [
        Path("C:/Windows/Fonts/NotoSansSC-VF.ttf"),
        Path("C:/Windows/Fonts/simfang.ttf"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsunb.ttf"),
    ]
    for font_path in font_candidates:
        if font_path.exists():
            pdfmetrics.registerFont(TTFont("ContractChinese", str(font_path)))
            return "ContractChinese"
    return "Helvetica"


def write_pdf(path: Path):
    font_name = register_pdf_font()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ContractTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=18,
        leading=24,
        alignment=1,
        spaceAfter=4,
        textColor="#0B2545",
    )
    subtitle_style = ParagraphStyle(
        "ContractSubtitle",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=10,
        leading=14,
        alignment=1,
        spaceAfter=14,
        textColor="#555555",
    )
    h1_style = ParagraphStyle(
        "ContractHeading1",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=15,
        leading=20,
        spaceBefore=12,
        spaceAfter=8,
        textColor="#2E74B5",
    )
    h2_style = ParagraphStyle(
        "ContractHeading2",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=12,
        leading=16,
        spaceBefore=8,
        spaceAfter=4,
        textColor="#2E74B5",
    )
    body_style = ParagraphStyle(
        "ContractBody",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=10.5,
        leading=16,
        spaceAfter=5,
        firstLineIndent=0,
    )

    story = [
        Paragraph(TITLE, title_style),
        Paragraph(SUBTITLE, subtitle_style),
        Paragraph("合同主数据", h1_style),
    ]
    for line in CONTRACT_LINES:
        story.append(Paragraph(line, body_style))
    story.extend([Spacer(1, 6), Paragraph("合同条款", h1_style)])
    for title_text, body in CLAUSES:
        story.append(Paragraph(title_text, h2_style))
        story.append(Paragraph(body, body_style))
    story.extend([Spacer(1, 6), Paragraph("签署页", h1_style)])
    for line in SIGNATURE_LINES:
        story.append(Paragraph(line, body_style))

    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=TITLE,
        author="contract-agent",
    )
    doc.build(story)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    write_txt(OUTPUT_DIR / f"{BASENAME}.txt")
    write_docx(OUTPUT_DIR / f"{BASENAME}.docx")
    write_pdf(OUTPUT_DIR / f"{BASENAME}.pdf")
    print(f"Generated files in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
